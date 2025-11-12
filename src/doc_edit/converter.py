"""
Document converter for PDF and Word documents to LaTeX.
Supports tagged PDFs with enhanced structure extraction.
"""

import os
import io
import re
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import pdfplumber
from docx import Document as DocxDocument
from docx.shared import Inches
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentConverter:
    """Converts various document formats to LaTeX with enhanced tagged document support."""

    def __init__(self):
        """Initialize the document converter."""
        # COMMENTED OUT: PDF/DOCX conversion disabled - only .tex files supported
        # self.supported_formats = ['.pdf', '.docx', '.doc', '.tex']
        self.supported_formats = ['.tex']  # Only LaTeX files for now
        
    def is_tagged_pdf(self, file_path: str) -> bool:
        """
        Check if a PDF is tagged (has accessibility structure).
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            True if PDF is tagged, False otherwise
        """
        try:
            doc = fitz.open(file_path)
            # Check for the existence of the '/StructTreeRoot' in the document's catalog
            catalog = doc.xref_object(doc.pdf_catalog(), compressed=False)
            is_tagged = isinstance(catalog, dict) and '/StructTreeRoot' in catalog
            
            # Also check MarkInfo
            mark_info = doc.markinfo
            has_mark_info = mark_info is not None and mark_info.get('Marked', False)
            
            doc.close()
            return is_tagged or has_mark_info
        except Exception as e:
            logger.warning(f"Error checking if PDF is tagged: {e}")
            return False

    def convert_to_latex(self, file_path: str) -> Dict[str, Any]:
        """
        Convert a document file to LaTeX.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with latex_code and metadata
        """
        file_extension = Path(file_path).suffix.lower()

        # COMMENTED OUT: PDF conversion - provide LaTeX files directly instead
        # if file_extension == '.pdf':
        #     return self._pdf_to_latex(file_path)
        if file_extension == '.pdf':
            raise ValueError("PDF conversion is currently disabled. Please provide a .tex file directly.")
        elif file_extension in ['.docx', '.doc']:
            return self._docx_to_latex(file_path)
        elif file_extension == '.tex':
            return self._tex_to_latex(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _pdf_to_latex(self, pdf_path: str) -> Dict[str, Any]:
        """
        Convert PDF to LaTeX using enhanced structure extraction for tagged PDFs.
        Currently only processes tagged PDFs - non-tagged PDFs are rejected.
        """
        try:
            is_tagged = self.is_tagged_pdf(pdf_path)
            logger.info(f"Processing {'tagged' if is_tagged else 'non-tagged'} PDF: {pdf_path}")
            
            if is_tagged:
                return self._tagged_pdf_to_latex(pdf_path)
            else:
                # FOR NOW: Only process tagged PDFs, reject non-tagged ones
                print("❌ NON-TAGGED DOCUMENT: This document does not have accessibility tags.")
                print("   Currently, only tagged PDFs are supported for processing.")
                print("   Please use an accessibility-tagged PDF instead.")
                raise ValueError("Non-tagged PDFs are currently not supported. Only tagged PDFs can be processed.")
                
                # COMMENTED OUT: Non-tagged PDF processing code (kept for future use)
                # return self._standard_pdf_to_latex(pdf_path)

        except ValueError:
            # Re-raise ValueError (our custom non-tagged PDF error)
            raise
        except Exception as e:
            raise RuntimeError(f"Failed to convert PDF to LaTeX: {str(e)}")
    
    def _tagged_pdf_to_latex(self, pdf_path: str) -> Dict[str, Any]:
        """
        Convert tagged PDF to LaTeX leveraging accessibility structure.
        Uses semantic tags when available, falls back to enhanced text extraction.
        """
        doc = fitz.open(pdf_path)
        metadata = {
            "source_format": "tagged_pdf",
            "pages": len(doc),
            "has_structure": True,
            "elements_found": []
        }
        
        try:
            # Try to use accessibility structure first
            latex_code = self._extract_from_accessibility_tags(doc)
            
            if not latex_code:
                # Fallback to enhanced text extraction
                logger.info("Accessibility structure not available, using enhanced text extraction")
                structured_content = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text_dict = page.get_text("dict")
                    page_content = self._extract_structured_content(text_dict, page_num)
                    structured_content.extend(page_content)
                
                latex_code = self._structured_content_to_latex(structured_content)
                element_types = set(item.get('type', 'text') for item in structured_content)
                metadata['elements_found'] = list(element_types)
            else:
                metadata['extraction_method'] = 'accessibility_tags'
                
        finally:
            doc.close()
            
        return {
            "latex_code": latex_code,
            "metadata": metadata
        }
    
    def _extract_from_accessibility_tags(self, doc) -> Optional[str]:
        """
        Extract content using PDF accessibility/structure tags.
        Returns LaTeX code if successful, None if tags not available.
        """
        try:
            # For now, this is a placeholder for future implementation
            # PyMuPDF doesn't have direct structure tree access in the current version
            # We would need to implement PDF structure tree parsing
            
            # Check if we can access structure information
            catalog = doc.xref_object(doc.pdf_catalog())
            if not isinstance(catalog, dict) or '/StructTreeRoot' not in catalog:
                return None
                
            # TODO: Implement structure tree traversal
            # This would involve:
            # 1. Parse the structure tree from the PDF
            # 2. Map structure elements to LaTeX equivalents
            # 3. Extract content in proper semantic order
            
            logger.info("Structure tree detected but parsing not yet implemented")
            return None
            
        except Exception as e:
            logger.warning(f"Could not extract from accessibility tags: {e}")
            return None
    
    # COMMENTED OUT: Standard PDF processing (for non-tagged PDFs)
    # This functionality is temporarily disabled - only tagged PDFs are currently supported
    # Keeping the code for future implementation when non-tagged PDF support is needed
    
    # def _standard_pdf_to_latex(self, pdf_path: str) -> Dict[str, Any]:
    #     """
    #     Convert standard PDF to LaTeX using basic text extraction.
    #     CURRENTLY DISABLED - Only tagged PDFs are supported.
    #     """
    #     doc = fitz.open(pdf_path)
    #     text_content = []
    #     
    #     try:
    #         for page_num in range(len(doc)):
    #             page = doc.load_page(page_num)
    #             text = page.get_text()
    #             text_content.append(text)
    #     finally:
    #         doc.close()
    # 
    #     # Basic LaTeX structure
    #     latex_code = self._text_to_basic_latex('\n'.join(text_content))
    # 
    #     return {
    #         "latex_code": latex_code,
    #         "metadata": {
    #             "source_format": "pdf",
    #             "pages": len(text_content),
    #             "has_structure": False
    #         }
    #     }
    
    def _extract_structured_content(self, text_dict: Dict, page_num: int) -> List[Dict[str, Any]]:
        """
        Extract structured content from PDF text dictionary.
        Groups spans within lines to preserve coherent text blocks and mathematical expressions.
        """
        content = []
        
        for block in text_dict.get("blocks", []):
            if "lines" in block:  # Text block
                for line in block["lines"]:
                    # Group all spans in the line together to preserve coherent text
                    line_text_parts = []
                    line_font_sizes = []
                    line_font_flags = []
                    line_font_names = []
                    line_bboxes = []
                    
                    for span in line["spans"]:
                        span_text = span["text"]
                        if span_text.strip():  # Only include non-empty spans
                            line_text_parts.append(span_text)
                            line_font_sizes.append(span.get("size", 12))
                            line_font_flags.append(span.get("flags", 0))
                            line_font_names.append(span.get("font", ""))
                            line_bboxes.append(span.get("bbox", []))
                    
                    if not line_text_parts:
                        continue
                    
                    # Combine all spans in the line into coherent text
                    full_line_text = "".join(line_text_parts).strip()
                    
                    if not full_line_text:
                        continue
                    
                    # Use the dominant formatting characteristics of the line
                    avg_font_size = sum(line_font_sizes) / len(line_font_sizes) if line_font_sizes else 12
                    dominant_flags = max(set(line_font_flags), key=line_font_flags.count) if line_font_flags else 0
                    dominant_font = max(set(line_font_names), key=line_font_names.count) if line_font_names else ""
                    
                    # Determine content type based on the full line
                    content_type = self._classify_text_element(full_line_text, avg_font_size, dominant_flags, dominant_font)
                    
                    content.append({
                        "type": content_type,
                        "text": full_line_text,
                        "page": page_num,
                        "font_size": avg_font_size,
                        "is_bold": bool(dominant_flags & 2**4),
                        "is_italic": bool(dominant_flags & 2**1),
                        "font_name": dominant_font,
                        "bbox": line_bboxes[0] if line_bboxes else []
                    })
        
        return content
    
    def _classify_text_element(self, text: str, font_size: float, font_flags: int, font_name: str) -> str:
        """
        Classify text element type based on formatting and content.
        Enhanced to better detect mathematical content and document structure.
        """
        # Check for mathematical expressions first (they should be preserved as-is)
        math_patterns = [
            r'[E=mc²]',  # E=mc²
            r'[F=ma]',   # F=ma
            r'[PV=nRT]', # Ideal gas law
            r'[\w\s]*=[\w\s]*[\+\-\*/\^²³¹⁰¹²³⁴⁵⁶⁷⁸⁹\(\)\[\]]+',  # General equations
            r'[∂∇∫∑∏√±≤≥≠≈∞π∆Ω∈∉⊂⊃∪∩]',  # Mathematical symbols
            r'[α-ωΑ-Ω]',  # Greek letters
            r'[\w]*\^\{.*\}',  # LaTeX-style superscripts
            r'[\w]*_\{.*\}',   # LaTeX-style subscripts
        ]
        
        for pattern in math_patterns:
            if re.search(pattern, text):
                return "math_expression"
        
        # Check for headings based on font size and formatting
        if font_size > 16:
            return "title"
        elif font_size > 14:
            return "section"
        elif font_size > 12:
            return "subsection"
        
        # Check for special patterns
        if re.match(r'^\d+\.?\s+', text):  # Numbered items
            return "enumeration"
        elif re.match(r'^[•\-\*]\s+', text):  # Bullet points
            return "itemize"
        elif text.isupper() and len(text) < 100:  # All caps short text
            return "heading"
        elif bool(font_flags & 2**4):  # Bold text
            return "emphasis"
        
        return "paragraph"
    
    def _structured_content_to_latex(self, content: List[Dict[str, Any]]) -> str:
        """
        Convert structured content to LaTeX with proper formatting.
        Enhanced for tagged PDF content with better semantic understanding.
        """
        latex_parts = [
            r"\documentclass[12pt]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage[T1]{fontenc}",
            r"\usepackage{lmodern}",
            r"\usepackage{geometry}",
            r"\usepackage{xcolor}",
            r"\usepackage{amsmath}",
            r"\usepackage{amsfonts}",
            r"\usepackage{amssymb}",
            r"\usepackage{graphicx}",
            r"\usepackage{enumitem}",
            r"\usepackage{hyperref}",
            r"\geometry{margin=1in}",
            "",
        ]
        
        # Try to extract title, author, date from content
        title_found = False
        for item in content:
            if item["type"] == "title" and not title_found:
                latex_parts.append(f"\\title{{{self._escape_latex(item['text'])}}}")
                title_found = True
                break
        
        if not title_found:
            latex_parts.append(r"\title{Converted Document}")
            
        # Look for author and date in the first few items
        author_found = False
        date_found = False
        for item in content[:10]:  # Check first 10 items
            text = item["text"].strip()
            if not author_found and ("agent" in text.lower() or "author" in text.lower()) and len(text) < 100:
                latex_parts.append(f"\\author{{{self._escape_latex(text)}}}")
                author_found = True
            elif not date_found and any(month in text for month in ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December", "2024", "2025"]):
                latex_parts.append(f"\\date{{{self._escape_latex(text)}}}")
                date_found = True
        
        if not author_found:
            latex_parts.append(r"\author{Document Converter}")
        if not date_found:
            latex_parts.append(r"\date{\today}")
            
        latex_parts.extend([
            "",
            r"\begin{document}",
            r"\maketitle",
            ""
        ])
        
        current_list_type = None
        in_list = False
        
        for item in content:
            text = self._escape_latex(item["text"])
            content_type = item["type"]
            
            # Handle different content types
            if content_type == "title":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\section{{{text}}}")
                
            elif content_type == "section":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\subsection{{{text}}}")
                
            elif content_type == "subsection":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\subsubsection{{{text}}}")
                
            elif content_type == "heading":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\paragraph{{{text}}}")
                
            elif content_type in ["enumeration", "itemize"]:
                list_type = "enumerate" if content_type == "enumeration" else "itemize"
                
                if not in_list or current_list_type != list_type:
                    if in_list:
                        latex_parts.append(f"\\end{{{current_list_type}}}")
                    latex_parts.append(f"\\begin{{{list_type}}}")
                    current_list_type = list_type
                    in_list = True
                
                # Clean up list item text
                clean_text = re.sub(r'^[\d\w]*\.?\s*[•\-\*]?\s*', '', text)
                latex_parts.append(f"\\item {clean_text}")
                
            elif content_type == "emphasis":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\textbf{{{text}}}")
                
            elif content_type == "math_expression":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                # Try to format as display math if it looks like a standalone equation
                if self._is_display_math(text):
                    latex_parts.append(f"\\[{self._convert_to_latex_math(text)}\\]")
                else:
                    # Inline math
                    latex_parts.append(f"${self._convert_to_latex_math(text)}$")
                
            else:  # paragraph
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                if text and len(text.strip()) > 0:
                    latex_parts.append(f"{text}")
                    latex_parts.append("")  # Add blank line after paragraphs
        
        # Close any open lists
        if in_list:
            latex_parts.append(f"\\end{{{current_list_type}}}")
        
        latex_parts.append("")
        latex_parts.append(r"\end{document}")
        
        return "\n".join(latex_parts)
    
    def _is_display_math(self, text: str) -> bool:
        """
        Determine if mathematical expression should be displayed as block math.
        """
        # Check if it's a standalone equation (contains = and looks substantial)
        if '=' in text and len(text.strip()) > 3:
            return True
        # Check for complex mathematical expressions
        complex_symbols = ['∫', '∑', '∏', '∂', '√', 'lim', 'frac']
        if any(symbol in text for symbol in complex_symbols):
            return True
        return False
    
    def _convert_to_latex_math(self, text: str) -> str:
        """
        Convert mathematical text to proper LaTeX math notation.
        """
        # Basic conversions for common mathematical expressions
        conversions = {
            '²': '^2',
            '³': '^3',
            '¹': '^1',
            '⁰': '^0',
            '⁴': '^4',
            '⁵': '^5',
            '⁶': '^6',
            '⁷': '^7',
            '⁸': '^8',
            '⁹': '^9',
            '±': '\\pm',
            '≤': '\\leq',
            '≥': '\\geq',
            '≠': '\\neq',
            '≈': '\\approx',
            '∞': '\\infty',
            'π': '\\pi',
            '∆': '\\Delta',
            'Ω': '\\Omega',
            '∈': '\\in',
            '∉': '\\notin',
            '⊂': '\\subset',
            '⊃': '\\supset',
            '∪': '\\cup',
            '∩': '\\cap',
            '∂': '\\partial',
            '∇': '\\nabla',
            '∫': '\\int',
            '∑': '\\sum',
            '∏': '\\prod',
            '√': '\\sqrt',
            'α': '\\alpha',
            'β': '\\beta',
            'γ': '\\gamma',
            'δ': '\\delta',
            'ε': '\\epsilon',
            'ζ': '\\zeta',
            'η': '\\eta',
            'θ': '\\theta',
            'ι': '\\iota',
            'κ': '\\kappa',
            'λ': '\\lambda',
            'μ': '\\mu',
            'ν': '\\nu',
            'ξ': '\\xi',
            'ο': '\\omicron',
            'π': '\\pi',
            'ρ': '\\rho',
            'σ': '\\sigma',
            'τ': '\\tau',
            'υ': '\\upsilon',
            'φ': '\\phi',
            'χ': '\\chi',
            'ψ': '\\psi',
            'ω': '\\omega',
            'ℏ': '\\hbar',
            'ψ': '\\psi',
        }
        
        result = text
        for unicode_char, latex_cmd in conversions.items():
            result = result.replace(unicode_char, latex_cmd)
        
        return result
    
    def _escape_latex(self, text: str) -> str:
        """
        Escape special LaTeX characters in text.
        
        Args:
            text: Text to escape
            
        Returns:
            LaTeX-escaped text
        """
        # Define LaTeX special characters and their escapes
        latex_special_chars = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\textasciicircum{}',
            '\\': r'\textbackslash{}'
        }
        
        # Replace special characters
        escaped_text = text
        for char, escape in latex_special_chars.items():
            escaped_text = escaped_text.replace(char, escape)
            
        return escaped_text

    def _docx_to_latex(self, docx_path: str) -> Dict[str, Any]:
        """
        Convert Word document to LaTeX with enhanced structure preservation.
        """
        try:
            doc = DocxDocument(docx_path)
            structured_content = []
            
            # Extract structured content from Word document
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                    
                # Analyze paragraph style and formatting
                style_name = paragraph.style.name
                text = paragraph.text.strip()
                
                # Determine content type
                if style_name.startswith('Heading'):
                    level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                    content_type = "title" if level == 1 else "section" if level == 2 else "subsection"
                elif style_name == 'List Paragraph' or text.startswith(('•', '-', '*')):
                    content_type = "itemize"
                elif re.match(r'^\d+\.', text):
                    content_type = "enumeration"
                elif style_name == 'Title':
                    content_type = "title"
                else:
                    content_type = "paragraph"
                
                # Check for formatting
                is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
                is_italic = any(run.italic for run in paragraph.runs if run.italic is not None)
                
                structured_content.append({
                    "type": content_type,
                    "text": text,
                    "style": style_name,
                    "is_bold": is_bold,
                    "is_italic": is_italic
                })
            
            # Extract tables
            tables_info = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_info.append(table_data)
            
            # Convert to LaTeX
            latex_code = self._docx_content_to_latex(structured_content, tables_info)
            
            # Extract metadata
            headings = [(item["type"], item["text"]) for item in structured_content 
                       if item["type"] in ["title", "section", "subsection"]]
            
            return {
                "latex_code": latex_code,
                "metadata": {
                    "source_format": "docx",
                    "headings": headings,
                    "has_tables": len(tables_info) > 0,
                    "table_count": len(tables_info),
                    "elements_found": list(set(item["type"] for item in structured_content))
                }
            }

        except Exception as e:
            raise RuntimeError(f"Failed to convert Word to LaTeX: {str(e)}")
    
    def _docx_content_to_latex(self, content: List[Dict[str, Any]], tables: List[List[List[str]]]) -> str:
        """
        Convert Word document content to LaTeX with proper formatting.
        """
        latex_parts = [
            r"\documentclass[12pt]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage[T1]{fontenc}",
            r"\usepackage{lmodern}",
            r"\usepackage{geometry}",
            r"\usepackage{xcolor}",
            r"\usepackage{amsmath}",
            r"\usepackage{graphicx}",
            r"\usepackage{enumitem}",
            r"\usepackage{array}",
            r"\usepackage{longtable}",
            r"\geometry{margin=1in}",
            "",
            r"\title{Converted Document}",
            r"\author{Document Converter}",
            r"\date{\today}",
            "",
            r"\begin{document}",
            r"\maketitle",
            ""
        ]
        
        current_list_type = None
        in_list = False
        table_index = 0
        
        for item in content:
            text = self._escape_latex(item["text"])
            content_type = item["type"]
            is_bold = item.get("is_bold", False)
            is_italic = item.get("is_italic", False)
            
            # Apply formatting
            if is_bold and is_italic:
                text = f"\\textbf{{\\textit{{{text}}}}}"
            elif is_bold:
                text = f"\\textbf{{{text}}}"
            elif is_italic:
                text = f"\\textit{{{text}}}"
            
            # Handle different content types
            if content_type == "title":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\section{{{text}}}")
                
            elif content_type == "section":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\subsection{{{text}}}")
                
            elif content_type == "subsection":
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                latex_parts.append(f"\\subsubsection{{{text}}}")
                
            elif content_type in ["enumeration", "itemize"]:
                list_type = "enumerate" if content_type == "enumeration" else "itemize"
                
                if not in_list or current_list_type != list_type:
                    if in_list:
                        latex_parts.append(f"\\end{{{current_list_type}}}")
                    latex_parts.append(f"\\begin{{{list_type}}}")
                    current_list_type = list_type
                    in_list = True
                
                # Clean up list item text
                clean_text = re.sub(r'^[\d\w]*\.?\s*[•\-\*]?\s*', '', text)
                latex_parts.append(f"\\item {clean_text}")
                
            else:  # paragraph
                if in_list:
                    latex_parts.append(f"\\end{{{current_list_type}}}")
                    in_list = False
                if text and len(text.strip()) > 0:
                    latex_parts.append(f"{text}")
                    latex_parts.append("")  # Add blank line after paragraphs
        
        # Close any open lists
        if in_list:
            latex_parts.append(f"\\end{{{current_list_type}}}")
        
        # Add tables
        for i, table in enumerate(tables):
            if table and len(table) > 0:
                latex_parts.append("")
                latex_parts.append(f"\\begin{{table}}[h]")
                latex_parts.append(f"\\centering")
                latex_parts.append(f"\\caption{{Table {i+1}}}")
                
                # Determine column specification
                num_cols = len(table[0]) if table else 1
                col_spec = "|" + "c|" * num_cols
                
                latex_parts.append(f"\\begin{{tabular}}{{{col_spec}}}")
                latex_parts.append("\\hline")
                
                for j, row in enumerate(table):
                    escaped_row = [self._escape_latex(cell) for cell in row]
                    latex_parts.append(" & ".join(escaped_row) + " \\\\")
                    latex_parts.append("\\hline")
                
                latex_parts.append("\\end{tabular}")
                latex_parts.append("\\end{table}")
        
        latex_parts.append("")
        latex_parts.append(r"\end{document}")
        
        return "\n".join(latex_parts)

    # COMMENTED OUT: Basic text to LaTeX conversion (used only for non-tagged PDFs)
    # This functionality is temporarily disabled along with non-tagged PDF support
    # Keeping the code for future implementation
    
    # def _text_to_basic_latex(self, text: str) -> str:
    #     """
    #     Convert plain text to basic LaTeX structure.
    #     CURRENTLY DISABLED - Only used for non-tagged PDFs.
    #     """
    #     # Basic LaTeX template
    #     latex_template = """\\documentclass[12pt]{article}
    # \\usepackage[utf8]{inputenc}
    # \\usepackage[T1]{fontenc}
    # \\usepackage{lmodern}
    # \\usepackage{geometry}
    # \\geometry{margin=1in}
    # 
    # \\title{Converted Document}
    # \\author{Document Converter}
    # \\date{\\today}
    # 
    # \\begin{document}
    # 
    # \\maketitle
    # 
    # \\section{Content}
    # 
    # {text_content}
    # 
    # \\end{document}"""
    # 
    #     # Basic text processing using the escape function
    #     processed_text = self._escape_latex(text)
    # 
    #     # Convert line breaks to paragraphs
    #     paragraphs = processed_text.split('\n\n')
    #     formatted_paragraphs = []
    # 
    #     for para in paragraphs:
    #         if para.strip():
    #             # Check if it looks like a heading
    #             if len(para.strip()) < 100 and not para.endswith('.'):
    #                 formatted_paragraphs.append(f'\\section{{{para.strip()}}}')
    #             else:
    #                 formatted_paragraphs.append(para.strip())
    # 
    #     final_text = '\n\n'.join(formatted_paragraphs)
    # 
    #     return latex_template.format(text_content=final_text)

    def _tex_to_latex(self, tex_path: str) -> Dict[str, Any]:
        """
        Read an existing LaTeX file.
        """
        try:
            with open(tex_path, 'r', encoding='utf-8') as f:
                latex_code = f.read()

            return {
                "latex_code": latex_code,
                "metadata": {
                    "source_format": "tex",
                    "source_file": tex_path
                }
            }

        except Exception as e:
            raise RuntimeError(f"Failed to read LaTeX file: {str(e)}")